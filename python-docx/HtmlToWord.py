from docx import Document
from docx.shared import Pt, RGBColor
from  docx.oxml.ns import  qn
from docx.shared import Inches
import requests
import os
import shutil

#桌面路径
kDesktopPath = '/Users/yimiaotong/Desktop/Words'

def questionnaireToWord(jsonData,fileName):
    document = Document()

    common_style = commonStytle(document)
    title_style = titleStytle(document)
    answer_style = answerStytle(document)

    for dic in jsonData:
        answers = dic['answer']
        title = dic['title']
        type = dic['type']
        document.add_paragraph(title, style = title_style)

        rowNum = ['A','B','C','D','E','F']
        if type == '2':
           document.add_paragraph(answers,style = answer_style)
        else:
            selects = dic['select']
            for i in range(len(answers)):
                answer = answers[i]
                if answer in selects:
                    document.add_paragraph(rowNum[i] + ':' + answer, style = answer_style)
                else:
                    document.add_paragraph(rowNum[i] + ':' + answer, style = common_style)

    try:
        # 保存文件
        err = document.save(fileName)
        if(err):
            print('转成 word 失败')
        else:
            print('转成 word 成功')

    except:
        print('转成 word 失败')

def commonStytle(document,styleStr='commonStytle'):
    style = document.styles.add_style(styleStr, 1)
    # 设置字体尺寸
    style.font.size = Pt(14)
    # 设置字体颜色
    style.font.color.rgb = RGBColor(0x3f, 0x2c, 0x36)

    # 居中文本
    # style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # # 设置中文字体
    # style.font.name = '微软雅黑'
    # style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # 设置段落样式为宋体
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return style

def titleStytle(document):
    style = commonStytle(document,'titleStytle')
    # 设置字体尺寸
    style.font.size = Pt(20)
    return style

def answerStytle(document):
    style = commonStytle(document,'answerStytle')
    # 设置字体尺寸
    style.font.color.rgb = RGBColor(219, 71, 71)
    return style

def addImageInWord(imgs,fileName):
    document = Document(fileName)
    r = document.add_paragraph().add_run()

    img_name = 'saveImage.png'
    for img_str in imgs:
        # 保存图片至本地
        with open(img_name, 'wb')as f:
            response = requests.get(img_str).content
            f.write(response)
            f.close()
        r.add_picture(img_name)
        r.add_text('                      ')

    try:
        # 保存文件
        err = document.save(fileName)
        if(err):
            print('图片写入 word 失败')
        else:
            print('图片写入 word 成功')

    except:
        print('图片写入 word 失败')
    #如果图片存在就删除
    if os.path.exists(img_name):
        os.remove(img_name)
    else:
        print('no such file %s'%img_name)

def move_file(src_path, dst_path,fileName):
    print ('from : '+src_path)
    print ('to : '+dst_path)
    try:
        if not os.path.exists(dst_path):
            os.mkdir(dst_path)
        f_dst = os.path.join(dst_path,fileName)
        shutil.move(src_path, f_dst)

    except Exception as e:
        print('move_file ERROR: ',e)
        # traceback.print_exc()

def pythonWord(jsonData,fileName):
    #问卷 json 转 word
    # fileName = 'word.docx'
    questionnaireToWord(jsonData,fileName)

    #添加图片到 word
    imgs_arr = [
        "http://knowledge-profile.yimiaoquan100.com/9ffceaf965514b2499e39045a0979e96.jpg",
        "http://knowledge-profile.yimiaoquan100.com/9ffceaf965514b2499e39045a0979e96.jpg"
    ]
    addImageInWord(imgs_arr,fileName)

    #移动 word 到桌面
    fromPath = os.path.abspath(fileName)
    move_file(fromPath,kDesktopPath,fileName)

if __name__ == '__main__':

    #请修改全局 kDesktopPath 设置为本地路径
    jsonData = [
        {
            'title':'草帽团',
            'answer':['路飞','索罗','山治','乔巴'],
            'select':['路飞'],
            'type': '0'

        },{
            'title':'梦幻角色',
            'answer':['剑侠客','逍遥生','飞燕女','巫蛮儿','英女侠'],
            'select': ['逍遥生','英女侠'],
            'type':'1'
        },{

            'title': '填空题',
            'answer': '正月里采花无哟花采，二月间采花花哟正开，二月间采花花哟正开。三月里桃花红哟似海，四月间葡萄架哟上开，四月间葡萄架哟上开。',
            'type':'2'
        }
    ]
    for i in range(2):
        fileName = 'word'+str(i)+'.docx'
        pythonWord(jsonData,fileName)

